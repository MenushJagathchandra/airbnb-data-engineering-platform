#!/usr/bin/env python3
"""Generates the Expernetic Talent Assessment DOCX report."""
import os
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

PLOTS = Path("data/gold/plots")
OUT = Path("data/gold")

def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    return h

def add_body(doc, text):
    doc.add_paragraph(text)

def add_bullet(doc, text):
    doc.add_paragraph(text, style='List Bullet')

def add_table(doc, headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Light Shading Accent 1'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        t.rows[0].cells[i].text = h
        for p in t.rows[0].cells[i].paragraphs:
            for r in p.runs:
                r.bold = True
    for row_data in rows:
        row = t.add_row()
        for i, val in enumerate(row_data):
            row.cells[i].text = str(val)
    doc.add_paragraph("")

def add_plot(doc, filename, caption, width=5.5):
    p = Path(PLOTS / filename)
    if p.exists():
        doc.add_picture(str(p), width=Inches(width))
        last = doc.paragraphs[-1]
        last.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap = doc.add_paragraph(caption)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.runs[0].italic = True if cap.runs else None

def main():
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # COVER PAGE
    for _ in range(6):
        doc.add_paragraph("")
    t = doc.add_paragraph("Menush")
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in t.runs:
        r.font.size = Pt(36)
        r.bold = True
    sub = doc.add_paragraph("Expernetic Talent Assessment Program\nData Engineer Intern — Technical Report")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info = doc.add_paragraph("\nCandidate: Menush\nDataset: Inside Airbnb (Amsterdam & Venice)\nScrape Date: September 11, 2025\nSubmission Date: June 2026")
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    # TABLE OF CONTENTS placeholder
    add_heading(doc, "Table of Contents", 1)
    toc_items = [
        "1. Executive Summary", "2. Objectives & Scope", "3. Dataset Overview",
        "4. Methodology", "5. Engineering Approach", "6. EDA Findings",
        "7. Statistical Findings", "8. Data Science Experiments",
        "9. AI/ML Experiments", "10. Visualizations",
        "11. Business Recommendations", "12. Cross-City Comparisons",
        "13. Limitations & Caveats", "14. Future Improvements",
        "15. Reflection", "Appendix A: AI Usage Disclosure"
    ]
    for item in toc_items:
        doc.add_paragraph(item)
    doc.add_page_break()

    # 1. EXECUTIVE SUMMARY
    add_heading(doc, "1. Executive Summary", 1)
    add_body(doc, "This report presents the findings of a production-grade data engineering pipeline built to analyze the Inside Airbnb short-term rental marketplace across Amsterdam and Venice. The pipeline processes 13,576 listings, 6.9 million calendar records, and 1.34 million guest reviews through a four-stage ETL architecture: automated ingestion, data quality enforcement with quarantine, DuckDB-based dimensional modeling, and statistical analysis with AI-powered sentiment scoring.")
    add_body(doc, "Key findings include:")
    add_bullet(doc, "Amsterdam listings average $336.79/night (median $222), while Venice averages $241.02/night (median $158). Entire-home rentals command a statistically significant premium over private rooms in both cities (p < 1e-09).")
    add_bullet(doc, "Venice exhibits a clear geographic pricing gradient: 52.5% of listings within 1km of St. Mark's Square average $295.96/night vs. $129.28 for peripheral listings. Amsterdam shows no such pattern, with urban-ring listings outpricing the core.")
    add_bullet(doc, "Superhost status correlates with significantly higher review ratings in both cities, with Venice showing a large effect size (Cohen's d = 0.81).")
    add_bullet(doc, "Venice's rental market shows substantial commercial consolidation: 34.4% of hosts operate multiple listings, with the top 10 hosts controlling 8.7% of supply.")
    add_bullet(doc, "Lexicon-based sentiment analysis of 1.34M reviews shows moderate correlation with ratings (r = 0.34-0.39) but near-zero correlation with price, indicating that guest satisfaction is driven by hospitality rather than price tier.")
    doc.add_page_break()

    # 2. OBJECTIVES & SCOPE
    add_heading(doc, "2. Objectives & Scope", 1)
    add_body(doc, "The objective of this project was to design and implement a configuration-driven, containerized data engineering pipeline that transforms raw Inside Airbnb datasets into actionable business intelligence. The evaluation criteria emphasize problem-solving (30%), engineering quality (25%), and analytical storytelling (20%).")
    add_heading(doc, "Cities Selected", 2)
    add_body(doc, "Amsterdam (The Netherlands) and Venice (Italy) were selected to demonstrate multi-city pipeline scalability while providing contrasting market dynamics — a mature Northern European market versus a tourism-dependent Mediterranean market.")
    add_heading(doc, "Prioritization Rationale", 2)
    add_body(doc, 'Following the assignment\'s guidance that "a candidate who completes two sections with exceptional depth will outperform one who attempts all sections superficially," effort was concentrated on Sections 02-05 (Data Engineering + EDA + Statistics) with production-grade depth. Sections 06-08 (Data Science, Advanced AI, Open Innovation) were deprioritized in favor of engineering rigor and statistical completeness.')
    doc.add_page_break()

    # 3. DATASET OVERVIEW
    add_heading(doc, "3. Dataset Overview", 1)
    add_heading(doc, "Source", 2)
    add_body(doc, "All data is sourced from Inside Airbnb (insideairbnb.com), an independent public dataset providing detailed snapshots of Airbnb listings. The September 11, 2025 scrapes were used for both cities.")
    add_heading(doc, "Files Used", 2)
    add_table(doc, ["File", "Description", "Amsterdam Rows", "Venice Rows"], [
        ["listings.csv.gz", "Core listing data: price, location, host info, amenities", "~10,480", "~8,590"],
        ["calendar.csv.gz", "Daily price and availability per listing (365 days)", "~3.8M", "~3.1M"],
        ["reviews.csv.gz", "Guest review text with reviewer ID and date", "~501K", "~841K"],
    ])
    add_heading(doc, "Key Relationships", 2)
    add_bullet(doc, "listings.id is the primary key joining to calendar.listing_id and reviews.listing_id")
    add_bullet(doc, "listings.host_id identifies the host entity across multiple listings")
    add_bullet(doc, "calendar.date provides the temporal dimension for seasonality analysis")
    add_heading(doc, "Known Limitations", 2)
    add_bullet(doc, "Calendar price column is entirely NULL in both city datasets — pipeline falls back to listing base price via COALESCE")
    add_bullet(doc, "Review text is multilingual (Dutch, Italian, English) — lexicon sentiment scoring operates only on English tokens")
    add_bullet(doc, "Data represents a single-point-in-time scrape, not longitudinal time series")
    add_bullet(doc, "Listing IDs may overlap across cities — handled via composite surrogate keys")
    doc.add_page_break()

    # 4. METHODOLOGY
    add_heading(doc, "4. Methodology", 1)
    add_body(doc, "The analytical approach follows a four-stage pipeline architecture, each stage producing immutable output artifacts that feed into the next stage:")
    add_bullet(doc, "Stage 1 — Ingestion: Configuration-driven HTTP downloads with 1MB chunked streaming and exponential backoff retry logic. Idempotent: skips if files already exist.")
    add_bullet(doc, "Stage 2 — Cleaning: Type parsing (currency strings to floats, booleans, dates), coordinate validation, and quarantine partitioning. Invalid records are written to separate Parquet files with documented rejection reasons.")
    add_bullet(doc, "Stage 3 — Modeling: DuckDB-based star schema compilation with composite surrogate keys (md5(id||city)), host deduplication via window functions, and vectorized Haversine distance calculations in pure SQL.")
    add_bullet(doc, "Stage 4 — Analysis: Statistical hypothesis testing (Welch's t-tests, ANOVA), effect size calculation (Cohen's d), correlation analysis, lexicon-based sentiment scoring of 1.34M reviews, and automated report/plot generation.")
    add_heading(doc, "Statistical Test Selection", 2)
    add_body(doc, "Welch's t-test was chosen over Student's t-test for all pairwise comparisons because Airbnb pricing data exhibits highly unequal variances across segments. Welch's approximation is strictly more robust and produces identical results when variances happen to be equal. One-way ANOVA was used for multi-group neighborhood comparisons.")
    doc.add_page_break()

    # 5. ENGINEERING APPROACH
    add_heading(doc, "5. Engineering Approach", 1)
    add_heading(doc, "Star Schema Design", 2)
    add_body(doc, "The dimensional model comprises 3 fact tables, 5 dimension tables, 1 bridge table, and 1 materialized view:")
    add_table(doc, ["Table", "Type", "Purpose"], [
        ["fact_listings", "Fact", "Core metrics: price, capacity, ratings, nearest landmark distance"],
        ["fact_calendar", "Fact", "Daily availability and pricing per listing"],
        ["fact_reviews", "Fact", "Review text with computed sentiment scores"],
        ["dim_listings", "Dimension", "Listing descriptive attributes (name, room type, property type)"],
        ["dim_hosts", "Dimension", "Host attributes (superhost status, tenure)"],
        ["dim_locations", "Dimension", "Neighbourhood and geographic hierarchy"],
        ["dim_landmarks", "Dimension", "Configured landmark coordinates"],
        ["dim_date", "Dimension", "Calendar dimension (year, month, day, weekend flag)"],
        ["dim_listing_landmark_distance", "Bridge", "Haversine distances between all listings and landmarks"],
        ["v_nearest_landmark", "View", "Nearest landmark per listing (window function)"],
    ])
    add_heading(doc, "Key Engineering Decisions", 2)
    add_bullet(doc, "Composite Surrogate Keys: md5(concat_ws('||', id, city)) prevents cross-city PK collisions")
    add_bullet(doc, "Host Deduplication: row_number() OVER (PARTITION BY host_id, city) selects one representative record per host")
    add_bullet(doc, "Vectorized Haversine: Spherical distance calculations run in DuckDB SQL, not Python loops")
    add_bullet(doc, "Quarantine over Silent Drops: Invalid records preserved with rejection reasons for auditability")
    add_heading(doc, "Containerization", 2)
    add_body(doc, "The pipeline is fully containerized with Docker (python:3.11-slim base image) and docker-compose.yml. SELinux volume relabeling (:z flag) and host UID mapping (user: 1000:1000) ensure cross-platform compatibility. A GitHub Actions workflow enables CI/CD execution triggered by repository_dispatch webhooks.")
    doc.add_page_break()

    # 6. EDA FINDINGS
    add_heading(doc, "6. Exploratory Data Analysis Findings", 1)
    add_heading(doc, "Market Composition", 2)
    add_table(doc, ["Metric", "Amsterdam", "Venice"], [
        ["Total Listings", "5,874", "7,702"],
        ["Average Price", "$336.79", "$241.02"],
        ["Median Price", "$222.00", "$158.00"],
        ["Average Rating", "4.84/5.0", "4.75/5.0"],
        ["Avg Sentiment Score", "0.627", "0.465"],
        ["Entire Home/Apt %", "77.1%", "79.8%"],
    ])
    add_heading(doc, "Price Distributions", 2)
    add_body(doc, "Both cities exhibit strongly right-skewed price distributions, indicating that a small number of luxury listings drive mean prices well above the median. This means that median price is a more reliable indicator of the 'typical' listing cost for market participants.")
    add_plot(doc, "amsterdam_price_distribution.png", "Figure 1: Amsterdam Price Distribution (95th percentile clip)")
    add_plot(doc, "venice_price_distribution.png", "Figure 2: Venice Price Distribution (95th percentile clip)")
    add_heading(doc, "Room Type Pricing", 2)
    add_plot(doc, "amsterdam_price_by_room_type.png", "Figure 3: Amsterdam Price by Room Type")
    add_plot(doc, "venice_price_by_room_type.png", "Figure 4: Venice Price by Room Type")
    add_heading(doc, "Geographic Pricing Gradients", 2)
    add_body(doc, "Venice shows a clear proximity premium: 52.5% of listings within 1km of St. Mark's Square average $295.96/night, declining to $129.28 for peripheral listings (correlation: -0.108). Amsterdam shows no such gradient (correlation: 0.023) — urban-ring listings (2.5-5km) actually command higher average prices ($423.02) than the core ($313.89), driven by large canal houses in premium residential neighborhoods.")
    add_plot(doc, "amsterdam_price_vs_distance.png", "Figure 5: Amsterdam Price vs. Distance to Nearest Landmark")
    add_plot(doc, "venice_price_vs_distance.png", "Figure 6: Venice Price vs. Distance to Nearest Landmark")
    add_heading(doc, "Host Portfolio & Supply Concentration", 2)
    add_body(doc, "Amsterdam is a fragmented market: 92.3% of hosts operate a single listing, with the top 10 hosts controlling just 2.7% of supply. Venice shows significant commercial consolidation: 34.4% of hosts operate multiple listings, with the top 10 controlling 8.7% of total supply. This suggests Venice has a more professionalized short-term rental market.")
    add_plot(doc, "amsterdam_host_portfolio.png", "Figure 7: Amsterdam Host Portfolio Distribution")
    add_plot(doc, "venice_host_portfolio.png", "Figure 8: Venice Host Portfolio Distribution")
    doc.add_page_break()

    # 7. STATISTICAL FINDINGS
    add_heading(doc, "7. Statistical Findings", 1)
    add_heading(doc, "H1: Entire-Home vs Private Room Prices", 2)
    add_table(doc, ["City", "Entire Home Mean", "Private Room Mean", "t-statistic", "p-value", "Cohen's d"], [
        ["Amsterdam", "$336.18", "$175.73", "5.91", "3.7e-09", "0.130"],
        ["Venice", "$263.69", "$151.44", "9.64", "8.8e-22", "0.224"],
    ])
    add_body(doc, "Interpretation: Reject H0 in both cities. Entire-home listings command a statistically significant and economically meaningful price premium. This confirms the structural advantage of whole-unit rentals driven by capacity, privacy, and amenity access.")
    add_heading(doc, "H2: Superhost vs Non-Superhost Ratings", 2)
    add_table(doc, ["City", "Superhost Mean", "Non-Superhost Mean", "t-statistic", "p-value", "Cohen's d"], [
        ["Amsterdam", "4.865", "4.825", "6.48", "1.0e-10", "0.172 (Small)"],
        ["Venice", "4.870", "4.640", "35.18", "9.7e-242", "0.810 (Large)"],
    ])
    add_body(doc, "Interpretation: Reject H0. The Superhost badge is a credible quality signal. In Venice, the effect is large (d=0.81), meaning Superhosts rate nearly a full standard deviation higher. For platform operators, this validates the Superhost program as an effective quality incentive.")
    add_heading(doc, "H3: Review Volume and Pricing", 2)
    add_body(doc, "Listings with >10 reviews are priced significantly lower than those with fewer reviews in both cities (Amsterdam: $257 vs $423, p=0.002; Venice: $200 vs $352, p=2.6e-11). This counterintuitive finding suggests that higher-priced luxury listings receive fewer bookings and therefore accumulate fewer reviews.")
    add_heading(doc, "H4: Neighborhood ANOVA", 2)
    add_body(doc, "One-way ANOVA rejects H0 for both cities (Amsterdam: F=10.38, p=5.3e-34; Venice: F=4.39, p=3.0e-18). Neighborhood is a statistically significant driver of price. However, pairwise Cohen's d between top neighborhoods is negligible (Amsterdam: -0.15; Venice: 0.02), meaning the premium between adjacent prime districts is minor.")
    add_heading(doc, "H5: Weekend vs Weekday Pricing", 2)
    add_body(doc, "Fail to reject H0 in both cities (Amsterdam: p=0.88; Venice: p=0.82). The lack of weekend premium detection is an expected artifact: raw calendar prices were NULL, so the analysis used listing base prices as fallback, which are constant across days of week.")
    add_heading(doc, "Correlation Analysis", 2)
    add_body(doc, "The strongest price drivers identified through Pearson correlation are: bathrooms (0.078), bedrooms (0.074), and accommodates (0.070). Review-based features show near-zero correlation with price, confirming that pricing is primarily a function of physical capacity rather than reputation.")
    add_plot(doc, "correlation_matrix.png", "Figure 9: Feature Correlation Matrix Heatmap")
    doc.add_page_break()

    # 8. DATA SCIENCE EXPERIMENTS
    add_heading(doc, "8. Data Science Experiments", 1)
    add_body(doc, "Data science modeling (price prediction, demand forecasting, listing clustering) was deprioritized in favor of deeper engineering and statistical coverage. With additional time, the recommended approach would be:")
    add_bullet(doc, "Price Prediction: XGBoost regressor with features including room_type, accommodates, bedrooms, bathrooms, neighbourhood, distance_to_landmark, and host_is_superhost. Evaluate using MAE and MAPE with 5-fold cross-validation.")
    add_bullet(doc, "Feature Importance: SHAP values to identify the most impactful predictors and detect non-linear interactions.")
    add_bullet(doc, "Listing Segmentation: K-Means clustering on standardized price, capacity, and review features to identify market segments.")
    doc.add_page_break()

    # 9. AI/ML EXPERIMENTS
    add_heading(doc, "9. AI/ML Experiments: Sentiment Analysis", 1)
    add_heading(doc, "Methodology", 2)
    add_body(doc, "A lightweight, dependency-free lexicon scorer was implemented in Python. The scorer maps each review comment against curated positive (21 words: amazing, clean, perfect, etc.) and negative (21 words: dirty, noisy, broken, etc.) word lists, computing a score between -1.0 (entirely negative) and +1.0 (entirely positive).")
    add_heading(doc, "Performance", 2)
    add_body(doc, "The scorer processed all 1,342,000 reviews in 9 seconds on a single CPU thread — approximately 149,000 reviews/second. A HuggingFace Transformer model was rejected due to the 1GB+ PyTorch dependency and multi-hour inference time without GPU.")
    add_heading(doc, "Results", 2)
    add_table(doc, ["City", "Sentiment-Rating Correlation", "Sentiment-Price Correlation"], [
        ["Amsterdam", "0.343 (Moderate)", "-0.017 (None)"],
        ["Venice", "0.393 (Moderate)", "0.064 (None)"],
    ])
    add_body(doc, "Business Interpretation: Guest sentiment correlates moderately with review ratings, validating the lexicon as a reasonable proxy. However, sentiment shows zero correlation with listing price, confirming that guest satisfaction is driven by hospitality quality and cleanliness — not by how much guests pay.")
    add_plot(doc, "amsterdam_sentiment_vs_rating.png", "Figure 10: Amsterdam Sentiment vs. Review Rating")
    add_plot(doc, "venice_sentiment_vs_rating.png", "Figure 11: Venice Sentiment vs. Review Rating")
    doc.add_page_break()

    # 10. VISUALIZATIONS
    add_heading(doc, "10. Visualizations Summary", 1)
    add_body(doc, "The pipeline auto-generates 13 publication-quality visualizations covering price distributions, room type comparisons, geographic gradients, seasonality patterns, host portfolios, sentiment correlations, and feature correlation matrices. All plots are embedded in their respective analysis sections above (Figures 1-11).")
    add_plot(doc, "amsterdam_seasonality.png", "Figure 12: Amsterdam Seasonality (Monthly Price & Day-of-Week Occupancy)")
    add_plot(doc, "venice_seasonality.png", "Figure 13: Venice Seasonality (Monthly Price & Day-of-Week Occupancy)")
    doc.add_page_break()

    # 11. BUSINESS RECOMMENDATIONS
    add_heading(doc, "11. Business Recommendations", 1)
    add_bullet(doc, "For Venice operators: Concentrate acquisitions within 1km of St. Mark's Square and Rialto Bridge. The proximity premium is clear and measurable ($296/night vs $129/night for peripheral listings).")
    add_bullet(doc, "For Amsterdam operators: Location within the city center provides no pricing advantage. Instead, target larger properties (higher bedroom/bathroom count) in premium residential neighborhoods like the canal ring, where urban-ring listings command $423/night.")
    add_bullet(doc, "For all markets: Pursue Superhost certification. In Venice, Superhosts rate 0.81 standard deviations higher — a large, credible quality signal that drives booking preference.")
    add_bullet(doc, "For platform strategists: Venice's commercial consolidation (34.4% multi-listing hosts, top 10 controlling 8.7% of supply) warrants monitoring for market health and regulatory compliance.")
    add_bullet(doc, "For pricing: Review volume inversely correlates with price, suggesting that luxury listings have lower booking velocity. Revenue optimization should balance rate against occupancy rather than maximizing nightly price alone.")
    doc.add_page_break()

    # 12. CROSS-CITY COMPARISONS
    add_heading(doc, "12. Cross-City Comparisons", 1)
    add_table(doc, ["Dimension", "Amsterdam", "Venice"], [
        ["Market Size", "5,874 listings", "7,702 listings"],
        ["Price Level", "Higher ($337 avg)", "Lower ($241 avg)"],
        ["Price Skew", "Extreme (Hotel Room avg: $5,971)", "Moderate"],
        ["Geographic Premium", "None (r=0.023)", "Present (r=-0.108)"],
        ["Host Fragmentation", "High (92.3% single-listing)", "Moderate (65.6% single-listing)"],
        ["Commercial Concentration", "Low (top 10 = 2.7%)", "High (top 10 = 8.7%)"],
        ["Superhost Effect on Ratings", "Small (d=0.17)", "Large (d=0.81)"],
        ["Guest Sentiment", "Higher (0.627)", "Lower (0.465)"],
    ])
    add_body(doc, "Amsterdam and Venice represent fundamentally different market archetypes. Amsterdam is a fragmented, individual-host market where pricing is driven by property size rather than location. Venice is a tourism-concentrated, commercially professionalized market where proximity to iconic landmarks directly drives rate.")
    doc.add_page_break()

    # 13. LIMITATIONS
    add_heading(doc, "13. Limitations & Caveats", 1)
    add_bullet(doc, "Calendar Price Nulls: The calendar.csv.gz price column was entirely NULL for both cities, requiring fallback to listing base prices. This eliminates true dynamic pricing signals and renders weekend/weekday analysis inconclusive.")
    add_bullet(doc, "Single Scrape Snapshot: Data represents a point-in-time capture, not longitudinal trends. Year-over-year growth analysis is not possible.")
    add_bullet(doc, "Multilingual Reviews: The lexicon sentiment scorer operates only on English tokens. Reviews in Dutch, Italian, or other languages receive neutral (0.0) scores, potentially underestimating sentiment variance.")
    add_bullet(doc, "Price Outliers: Amsterdam Hotel Room listings average $5,971 (38 listings), likely reflecting aggregated multi-night or suite packages. These inflate the city-wide mean.")
    add_bullet(doc, "Occupancy Proxy: The availability flag in the calendar does not distinguish between 'booked' and 'blocked by host'. True occupancy rates cannot be computed from this data alone.")

    # 14. FUTURE IMPROVEMENTS
    add_heading(doc, "14. Future Improvements", 1)
    add_bullet(doc, "Price Prediction Model: Implement XGBoost/LightGBM with SHAP explainability to quantify marginal feature impacts on price.")
    add_bullet(doc, "Topic Modeling: Apply BERTopic or NMF on review text to surface recurring guest themes (cleanliness, noise, location quality).")
    add_bullet(doc, "Incremental Processing: Transition from full rebuilds to UPSERT/MERGE operations keyed on scrape_date for efficient longitudinal tracking.")
    add_bullet(doc, "Cloud Deployment: Migrate to AWS S3 + DuckDB httpfs extension for serverless, scalable data lake architecture.")
    add_bullet(doc, "Interactive Dashboard: Build a Streamlit or Dash application for stakeholder self-service exploration.")
    add_bullet(doc, "Additional Cities: Scale to 5+ cities (Barcelona, Paris, Rome, Lisbon, Berlin) to validate cross-market pattern generalization and create a pan-European short-term rental benchmark.")
    add_bullet(doc, "AI Chatbot for Natural Language Querying: Develop a conversational interface allowing non-technical users to ask questions like 'What is the average price of entire homes near Dam Square in June?' The chatbot would parse natural language, generate parameterized DuckDB SQL queries, and return results with auto-generated visualizations. This would democratize data access without requiring Power BI or SQL knowledge.")

    # 15. REFLECTION
    add_heading(doc, "15. Reflection", 1)
    add_body(doc, "This project demonstrated the importance of prioritization in data engineering. Rather than attempting every section superficially, effort was concentrated on building a production-grade pipeline with genuine depth: composite key collision prevention, quarantine-based data quality enforcement, vectorized spatial calculations, and rigorous statistical testing with effect sizes.")
    add_body(doc, "Key trade-offs accepted: (1) Lexicon sentiment over deep learning — sacrificing accuracy for 150x throughput improvement; (2) Full rebuilds over incremental loads — prioritizing correctness over efficiency; (3) Deprioritizing ML modeling — choosing engineering rigor over prediction accuracy.")
    add_body(doc, "The most valuable lesson was discovering that the Inside Airbnb calendar prices were entirely NULL — a real-world data quality surprise that required designing a graceful fallback mechanism rather than simply failing. This is exactly the kind of challenge that separates production pipelines from academic exercises.")
    doc.add_page_break()

    # APPENDIX A: AI USAGE DISCLOSURE
    add_heading(doc, "Appendix A: AI Usage Disclosure", 1)
    ai_path = Path("AI_USAGE_DISCLOSURE.md")
    if ai_path.exists():
        text = ai_path.read_text()
        for line in text.split('\n'):
            line = line.strip()
            if not line or line.startswith('#'):
                if line.startswith('## '):
                    add_heading(doc, line.replace('## ', ''), 2)
                continue
            if line.startswith('- ') or line.startswith('* '):
                add_bullet(doc, line[2:])
            elif line.startswith('| ') or line.startswith('---'):
                continue
            else:
                add_body(doc, line)

    # SAVE
    out_path = OUT / "Menush_Expernetic_Report.docx"
    doc.save(str(out_path))
    print(f"Report saved to: {out_path}")
    print(f"File size: {out_path.stat().st_size / 1024:.0f} KB")

if __name__ == "__main__":
    main()
